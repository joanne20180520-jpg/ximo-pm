"""Build Lark knowledge-base exports: HTML (copy-paste) and DOCX (import)."""
import re
import base64
import pathlib
import shutil
import sys
import importlib.util

_preview_path = pathlib.Path(__file__).with_name('build-manual-preview.py')
_spec = importlib.util.spec_from_file_location('build_manual_preview', _preview_path)
preview = importlib.util.module_from_spec(_spec)
assert _spec.loader
_spec.loader.exec_module(preview)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

FONT_NAME = 'PingFang TC'
CHROME_PATH = '/Applications/Google Chrome.app/Contents/MacOS/Google Chrome'

md_path = preview.md_path
base = preview.base
lines = preview.lines
text = preview.text
shots_src = preview.shots_src

kb_html = md_path.with_name('教育訓練操作手冊-知識庫.html')
kb_docx = md_path.with_name('教育訓練操作手冊-知識庫.docx')
kb_lark_md = md_path.with_name('教育訓練操作手冊-lark.md')

MANUAL_PUBLIC_BASE = 'https://ximo-pm-umber.vercel.app/manual/screenshots'

KB_BANNER = '''
<div class="kb-banner" contenteditable="false">
  <strong>貼到 Lark 知識庫</strong>
  <ol>
    <li><strong>方式 A（最建議）</strong>：在 Lark 文件點 <strong>⋯ → 導入 → Word</strong>，選 <code>教育訓練操作手冊-知識庫.docx</code></li>
    <li><strong>方式 B</strong>：Lark 文件 <strong>⋯ → 導入 → Markdown</strong>，選 <code>教育訓練操作手冊-lark.md</code></li>
    <li><strong>方式 C</strong>：Word 全選複製後貼到 Lark 文件</li>
  </ol>
  <p>貼上後可刪除上方這段說明。線上版連結：<a href="https://ximo-pm-umber.vercel.app/manual/">https://ximo-pm-umber.vercel.app/manual/</a></p>
</div>
'''

KB_EXTRA_STYLE = '''
.kb-banner{margin:0 0 24px;padding:16px 18px;background:#fff8e8;border:1px solid #e0a24a;border-radius:10px;font-size:13.5px;line-height:1.65}
.kb-banner strong{color:#5a4320}
.kb-banner ol{margin:8px 0 0;padding-left:1.3em}
.kb-banner li{margin:4px 0}
.kb-banner p{margin:10px 0 0;color:#6a5320}
.card.kb-stack{display:block}
.card.kb-stack .lr-img{margin-bottom:12px}
main.kb-copy{max-width:900px}
'''


def render_kb_html():
    preview.EMBED_MODE = 'remote_url'
    html_parts = []
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                html_parts.append(f'<div class="flow-line"><code>{preview.esc(chr(10).join(code_buf))}</code></div>')
                in_code = False
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img:
            alt, src = img.group(1), img.group(2)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].startswith('|'):
                rows = []
                while j < len(lines) and lines[j].startswith('|'):
                    rows.append(lines[j])
                    j += 1
                if 'flowchart' in src:
                    html_parts.append(f'<figure class="flowchart">{preview.img_tag(alt, src)}</figure>')
                    html_parts.append(f'<div class="table-wrap">{preview.parse_table(rows)}</div>')
                else:
                    html_parts.append(
                        f'<section class="card kb-stack">'
                        f'<div class="lr-img">{preview.img_tag(alt, src)}</div>'
                        f'<div class="lr-txt">{preview.parse_table(rows)}</div>'
                        f'</section>'
                    )
                i = j
                continue
            html_parts.append(f'<figure class="{preview.figure_class(src)}">{preview.img_tag(alt, src)}</figure>')
            i += 1
            continue
        if line.strip() == ':::flow':
            body = []
            i += 1
            while i < len(lines) and lines[i].strip() != ':::':
                body.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].strip() == ':::':
                i += 1
            html_parts.append(preview.parse_flow_cards(body))
            continue
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                rows.append(lines[i])
                i += 1
            html_parts.append(f'<div class="table-wrap">{preview.parse_table(rows)}</div>')
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith('# '):
            html_parts.append(f'<header class="doc-head"><h1>{preview.inline(line[2:])}</h1></header>')
        elif line.startswith('## '):
            html_parts.append(preview.heading_html('h2', line[3:]))
        elif line.startswith('### '):
            html_parts.append(preview.heading_html('h3', line[4:]))
        elif line.startswith('#### '):
            html_parts.append(preview.heading_html('h4', line[5:]))
        elif line.startswith('> '):
            html_parts.append(f'<blockquote>{preview.inline(line[2:])}</blockquote>')
        elif line.startswith('- '):
            items = []
            while i < len(lines) and lines[i].startswith('- '):
                items.append(lines[i][2:])
                i += 1
            html_parts.append('<ul>' + ''.join(f'<li>{preview.inline(x)}</li>' for x in items) + '</ul>')
            continue
        elif re.match(r'^\d+\. ', line):
            html_parts.append('<div class="checklist">')
            while i < len(lines) and re.match(r'^\d+\. ', lines[i]):
                m = re.match(r'^(\d+)\. (.*)$', lines[i])
                html_parts.append(
                    f'<div class="check-item"><span class="badge">{m.group(1)}</span>'
                    f'<div>{preview.inline(m.group(2))}</div></div>'
                )
                i += 1
            html_parts.append('</div>')
            continue
        elif line.strip() == '---':
            html_parts.append('<hr/>')
        elif line.strip() == '待更新':
            html_parts.append('<p class="pending">待更新</p>')
        else:
            html_parts.append(f'<p>{preview.inline(line)}</p>')
        i += 1
    return (
        f'<!DOCTYPE html><html lang="zh-Hant"><head><meta charset="utf-8">'
        f'<meta name="viewport" content="width=device-width, initial-scale=1">'
        f'<title>璽墨專案管理系統 · 操作手冊（知識庫版）</title>'
        f'<style>{preview.style}{KB_EXTRA_STYLE}</style></head>'
        f'<body>{KB_BANNER}<main class="kb-copy">{"".join(html_parts)}</main></body></html>'
    )


def render_flowchart_png():
    svg = shots_src / 'flowchart.svg'
    png = shots_src / 'flowchart.png'
    if not svg.exists() or not pathlib.Path(CHROME_PATH).exists():
        return
    import subprocess
    subprocess.run(
        [
            CHROME_PATH, '--headless=new', '--disable-gpu', '--hide-scrollbars',
            f'--screenshot={png}', '--window-size=920,1000', svg.resolve().as_uri(),
        ],
        check=True,
        capture_output=True,
    )


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def style_paragraph_font(paragraph, size=Pt(11)):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def picture_width(img_path):
    rel = None
    name = img_path.name
    for src in re.findall(r'!\[[^\]]*\]\(([^)]+)\)', text):
        if pathlib.Path(src).name == name or (name == 'flowchart.png' and 'flowchart' in src):
            rel = src
            break
    _, w, h = preview.img_meta(rel or f'./training-screenshots/{name}')
    if h and h <= 180:
        return Inches(3.0)
    if 'flowchart' in name:
        return Inches(5.6)
    if w and h and w / h > 1.4:
        return Inches(5.8)
    return Inches(4.8)


def strip_md(s):
    s = re.sub(r'\*\*([^*]+)\*\*', r'\1', s)
    s = re.sub(r'`([^`]+)`', r'\1', s)
    s = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', s)
    return s.strip()


def add_rich_text(paragraph, text_line, size=Pt(11)):
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text_line)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size)
            run.font.name = 'Menlo'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Menlo')
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_picture(doc, img_path):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=picture_width(img_path))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)


def resolve_image(src):
    src_path = (base / src.lstrip('./')).resolve()
    if not src_path.exists():
        return None
    if src_path.suffix.lower() == '.svg':
        png = src_path.with_suffix('.png')
        if png.exists():
            return png
        return None
    return src_path


def table_col_widths(headers):
    n = len(headers)
    if n == 2:
        return [Cm(1.4), Cm(14.6)]
    if n == 3 and headers[0] in ('#', '標號', '順序'):
        return [Cm(1.2), Cm(3.5), Cm(11.3)]
    if n == 3:
        return [Cm(1.2), Cm(3.0), Cm(11.8)]
    if n == 4:
        return [Cm(1.2), Cm(2.5), Cm(5.5), Cm(6.8)]
    if n == 5:
        return [Cm(1.0), Cm(2.2), Cm(4.5), Cm(3.5), Cm(4.8)]
    return [Cm(16 / n)] * n


def add_table(doc, rows):
    headers = [c.strip() for c in rows[0].strip('|').split('|')]
    body_rows = []
    for row in rows[1:]:
        cells = [c.strip() for c in row.strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c or '') for c in cells):
            continue
        body_rows.append(cells)
    if not body_rows:
        return
    table = doc.add_table(rows=1 + len(body_rows), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    widths = table_col_widths(headers)
    for ci, w in enumerate(widths):
        if ci < len(table.columns):
            table.columns[ci].width = w
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ''
        p = cell.paragraphs[0]
        add_rich_text(p, h, size=Pt(10.5))
        for run in p.runs:
            run.bold = True
    for ri, row in enumerate(body_rows, start=1):
        for ci, c in enumerate(row):
            if ci >= len(table.rows[ri].cells):
                continue
            cell = table.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ''
            p = cell.paragraphs[0]
            add_rich_text(p, c, size=Pt(10.5))
    spacer = doc.add_paragraph('')
    spacer.paragraph_format.space_after = Pt(10)


def add_flow_cards(doc, body_lines):
    for line in body_lines:
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split('｜')]
        if len(parts) < 4:
            continue
        num, title, action, result = parts[0], parts[1], parts[2], parts[3]
        is_note = len(parts) >= 5 and parts[4] == 'note'
        p = doc.add_paragraph()
        run = p.add_run(f'{num}｜{title}')
        set_run_font(run, size=Pt(11), bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph()
        add_rich_text(p, f'操作：{action}', size=Pt(10.5))
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph()
        if is_note:
            run = p.add_run('要請款時 ')
            set_run_font(run, size=Pt(10.5), bold=True)
            add_rich_text(p, result, size=Pt(10.5))
        else:
            add_rich_text(p, f'→ {result}', size=Pt(10.5))
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph('')


def build_docx():
    render_flowchart_png()
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    normal.font.size = Pt(11)
    for level in range(1, 4):
        style = doc.styles[f'Heading {level}']
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    i = 0
    while i < len(lines):
        line = lines[i]
        img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img:
            alt, src = img.group(1), img.group(2)
            img_path = resolve_image(src)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if img_path:
                add_picture(doc, img_path)
            elif src.endswith('.svg'):
                p = doc.add_paragraph(f'[流程圖：{alt}]')
                set_run_font(p.runs[0], italic=True)
            if j < len(lines) and lines[j].startswith('|'):
                rows = []
                while j < len(lines) and lines[j].startswith('|'):
                    rows.append(lines[j])
                    j += 1
                add_table(doc, rows)
                i = j
                continue
            i += 1
            continue
        if line.strip() == ':::flow':
            body = []
            i += 1
            while i < len(lines) and lines[i].strip() != ':::':
                body.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].strip() == ':::':
                i += 1
            add_flow_cards(doc, body)
            continue
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith('# '):
            h = doc.add_heading(strip_md(line[2:]), level=0)
            style_paragraph_font(h)
        elif line.startswith('## '):
            m = re.match(r'^(.*?)\s*\{#([a-zA-Z0-9_-]+)\}\s*$', line[3:])
            h = doc.add_heading(strip_md(m.group(1) if m else line[3:]), level=1)
            style_paragraph_font(h)
        elif line.startswith('### '):
            m = re.match(r'^(.*?)\s*\{#([a-zA-Z0-9_-]+)\}\s*$', line[4:])
            h = doc.add_heading(strip_md(m.group(1) if m else line[4:]), level=2)
            style_paragraph_font(h)
        elif line.startswith('#### '):
            h = doc.add_heading(strip_md(line[5:]), level=3)
            style_paragraph_font(h)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            add_rich_text(p, line[2:])
            p.paragraph_format.left_indent = Inches(0.25)
        elif line.startswith('- '):
            while i < len(lines) and lines[i].startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                add_rich_text(p, lines[i][2:])
                style_paragraph_font(p)
                i += 1
            continue
        elif re.match(r'^\[.+\]\(#', line):
            plain = re.sub(r'\[([^\]]+)\]\(#[^)]+\)', r'\1', line)
            p = doc.add_paragraph(plain)
            set_run_font(p.runs[0], size=Pt(10), color=RGBColor(0x6A, 0x65, 0x5C))
            p.paragraph_format.space_after = Pt(6)
        elif line.strip() == '---':
            p = doc.add_paragraph('')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif line.strip() == '待更新':
            p = doc.add_paragraph('待更新')
            set_run_font(p.runs[0], bold=True)
        else:
            p = doc.add_paragraph()
            add_rich_text(p, line)
        i += 1

    doc.add_paragraph('')
    p = doc.add_paragraph('線上版：')
    set_run_font(p.runs[0])
    run = p.add_run('https://ximo-pm-umber.vercel.app/manual/')
    set_run_font(run, color=RGBColor(0x2D, 0x6A, 0x4F))
    doc.save(kb_docx)


def lark_image_url(src):
    name = pathlib.Path(src.lstrip('./')).name
    if name == 'flowchart.svg':
        name = 'flowchart.png'
    return f'{MANUAL_PUBLIC_BASE}/{name}'


def strip_heading_anchor(text):
    return re.sub(r'\s*\{#[a-zA-Z0-9_-]+\}\s*$', '', text).strip()


def flow_cards_to_md(body_lines):
    rows = ['| 步驟 | 標題 | 操作 | 說明 |', '| --- | --- | --- | --- |']
    for line in body_lines:
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split('｜')]
        if len(parts) < 4:
            continue
        num, title, action, result = parts[0], parts[1], parts[2], parts[3]
        is_note = len(parts) >= 5 and parts[4] == 'note'
        note = result if is_note else f'→ {result}'
        rows.append(f'| {num} | {title} | {action} | {note} |')
    return '\n'.join(rows) + '\n'


def build_lark_markdown():
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == ':::flow':
            body = []
            i += 1
            while i < len(lines) and lines[i].strip() != ':::':
                body.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].strip() == ':::':
                i += 1
            out.append(flow_cards_to_md(body))
            continue
        img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img:
            alt, src = img.group(1), img.group(2)
            out.append(f'![{alt}]({lark_image_url(src)})\n')
            i += 1
            continue
        if line.startswith('# '):
            out.append(f'# {strip_heading_anchor(line[2:])}\n')
        elif line.startswith('## '):
            out.append(f'## {strip_heading_anchor(line[3:])}\n')
        elif line.startswith('### '):
            out.append(f'### {strip_heading_anchor(line[4:])}\n')
        elif line.startswith('#### '):
            out.append(f'#### {strip_heading_anchor(line[5:])}\n')
        elif re.match(r'^\[.+\]\(#', line):
            out.append(re.sub(r'\[([^\]]+)\]\(#[^)]+\)', r'\1', line) + '\n')
        elif line.strip() == '---':
            out.append('\n---\n\n')
        else:
            out.append(line + '\n')
        i += 1
    footer = (
        '\n---\n\n'
        '線上版（需登入）：https://ximo-pm-umber.vercel.app/manual/\n'
    )
    kb_lark_md.write_text(''.join(out).strip() + footer, encoding='utf-8')


if __name__ == '__main__':
    import subprocess
    subprocess.run([sys.executable, str(_preview_path)], check=True)
    kb_html.write_text(render_kb_html(), encoding='utf-8')
    print('built', kb_html)
    build_docx()
    print('built', kb_docx)
    build_lark_markdown()
    print('built', kb_lark_md)
